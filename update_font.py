import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def update_english_font(doc_path):
    doc = Document(doc_path)
    
    def apply_font_to_runs(runs):
        for run in runs:
            if not run.text:
                continue
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.first_child_found_in("w:rFonts")
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            
            # 仅修改 w:ascii 和 w:hAnsi 字体，以应用到英文字符和数字上
            # 这样就不会影响汉字所使用的 w:eastAsia 字体
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # 处理正文段落
    for p in doc.paragraphs:
        apply_font_to_runs(p.runs)
        
    # 处理表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_font_to_runs(p.runs)
                    
    # 处理页眉页脚
    for section in doc.sections:
        for p in section.header.paragraphs:
            apply_font_to_runs(p.runs)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        apply_font_to_runs(p.runs)
        for p in section.footer.paragraphs:
            apply_font_to_runs(p.runs)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        apply_font_to_runs(p.runs)
                    
    doc.save(doc_path)
    print("Success: 英文/数字 字体已被全部修改为 Times New Roman")

if __name__ == "__main__":
    doc_path = r"e:\桌面\zuoye\毕业设计\src\论文\论文1.3.docx"
    update_english_font(doc_path)

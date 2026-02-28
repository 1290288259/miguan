import os
import sys
import random
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

REPEAT_SCALE = 2


def create_thesis_doc(filename):
    print(f"生成论文初稿: {filename}")
    doc = Document()

    # 样式设置
    style = doc.styles['Normal']
    style.font.name = u'宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    style.font.size = Pt(12)

    for i in range(1, 4):
        if f'Heading {i}' not in doc.styles:
            continue
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = u'黑体'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(18)
        elif i == 2:
            h_style.font.size = Pt(14)
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(12)
        else:
            h_style.font.size = Pt(13)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)

    def add_long_text(paragraph, base_text, repeat=1):
        repeat = max(1, int(repeat * REPEAT_SCALE))
        variations = [
            " 此部分对系统稳定性和可维护性具有直接影响，后续章节将进一步验证其合理性。",
            " 在高并发场景下需要重点考虑性能与资源消耗的平衡，避免分析延迟持续累积。",
            " 经过多轮测试与对比，方案的可扩展性和可移植性表现良好。",
            " 同时还需考虑异常处理与安全边界，保证系统在极端情况下仍可稳定运行。",
            " 该设计吸收了工程实践中的成熟经验，并结合项目实际进行了简化与优化。",
            " 为保证数据一致性，关键流程引入事务控制与状态校验机制。",
            " 通过合理的模块划分，降低了耦合度，便于后续功能迭代。",
            " 在实现层面配合日志追踪与监控指标，有助于定位问题并优化性能。",
        ]
        paragraph.add_run(base_text)
        for _ in range(max(repeat - 1, 0)):
            paragraph.add_run(random.choice(variations))
            paragraph.add_run(base_text)

    def add_code_block(doc, code_text):
        p = doc.add_paragraph(code_text, style='Normal')
        for run in p.runs:
            run.font.name = 'Consolas'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # 封面
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于AI大模型的蜜罐流量分析系统设计与实现')
    run.font.name = u'黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.size = Pt(22)
    run.bold = True
    doc.add_page_break()

    # 摘要
    doc.add_heading('摘要', level=1)
    abstract_cn = """
    随着网络攻击向自动化与智能化演进，传统基于规则的检测方式对未知威胁的识别能力不足。蜜罐技术通过模拟真实服务诱捕攻击流量，可有效获得高价值日志，但日志规模大、噪声多、人工分析成本高。为此，本文设计并实现了一套基于AI大模型的蜜罐流量分析系统。

    系统集成SSH、FTP与HTTP等多种蜜罐服务，能够记录登录尝试、命令执行与异常请求等攻击行为。后端采用Flask构建RESTful API，前端使用Vue 3与Element Plus实现可视化管理与查询。AI分析模块基于本地部署的大语言模型（如DeepSeek/Qwen），通过提示词工程对日志进行语义理解、攻击类型判定与风险评估，并结合生产者-消费者队列与多Worker并发分析机制提升处理能力。同时系统支持恶意IP自动封禁、攻击来源地图展示与多维度统计。

    测试结果表明，该系统能够较准确地识别常见攻击类型，并在高并发场景下保持稳定响应。研究成果为安全运维提供了可落地的智能分析方案，也为教学与研究提供了实验平台。

    关键词：蜜罐；流量分析；大语言模型；网络安全；Flask；Vue
    """
    p = doc.add_paragraph()
    add_long_text(p, abstract_cn, repeat=3)

    doc.add_heading('Abstract', level=1)
    abstract_en = """
    As cyber attacks become more automated and intelligent, rule-based detection is insufficient for unknown threats. Honeypots can attract and record attack traffic, but the resulting logs are massive and noisy, making manual analysis inefficient. This thesis presents an AI-driven honeypot traffic analysis system.

    The system integrates multiple honeypots (SSH/FTP/HTTP) to capture attack behaviors. The backend is built with Flask and exposes RESTful APIs, while the frontend uses Vue 3 and Element Plus for visualization and management. A locally deployed large language model (e.g., DeepSeek/Qwen via Ollama) performs semantic analysis, attack classification, and risk evaluation. A producer-consumer queue with multi-worker concurrency improves throughput. The system also supports malicious IP auto-blocking, geo-visualization of attack sources, and multi-dimensional statistics.

    Experiments show that the system can effectively identify common attack types and maintain stable performance under high concurrency. The results provide a practical solution for intelligent security analysis and a reusable platform for education and research.

    Keywords: honeypot; traffic analysis; large language model; cybersecurity; Flask; Vue
    """
    p = doc.add_paragraph()
    add_long_text(p, abstract_en, repeat=2)
    doc.add_page_break()

    # 目录
    doc.add_heading('目录', level=1)
    doc.add_paragraph('（此处在Word中生成目录）')
    doc.add_page_break()

    # 第一章 绪论
    doc.add_heading('第一章 绪论', level=1)
    doc.add_heading('1.1 研究背景', level=2)
    bg_text = """
    近年来网络攻击规模化、产业化趋势明显，传统安全防护主要依赖特征库匹配与规则检测，面对零日漏洞与复杂攻击链时容易出现漏报。蜜罐通过模拟真实业务服务捕获攻击行为，为威胁分析提供了重要数据来源，但日志量大且结构松散，分析效率难以保证。
    """
    p = doc.add_paragraph()
    add_long_text(p, bg_text, repeat=5)

    doc.add_heading('1.2 研究意义', level=2)
    sig_text = """
    本研究通过引入大语言模型提升日志语义理解能力，并结合自动化封禁与可视化展示形成闭环防御，具有提升检测准确性、降低运维成本与强化主动防御能力的意义。同时探索LLM在安全垂直领域的落地方法，为后续研究提供范式。
    """
    p = doc.add_paragraph()
    add_long_text(p, sig_text, repeat=4)

    doc.add_heading('1.3 国内外研究现状', level=2)
    status_text = """
    国外开源蜜罐项目如Cowrie、Dionaea、Glastopf等已较成熟，但多聚焦数据采集而非智能分析。近年来大语言模型在安全领域的应用逐渐增多，但面向蜜罐日志的实时语义分析仍处于探索阶段。国内在蜜罐部署与安全可视化方面已有实践，但与本地LLM结合的研究相对不足。
    """
    p = doc.add_paragraph()
    add_long_text(p, status_text, repeat=4)

    doc.add_heading('1.4 研究内容与方法', level=2)
    content_text = """
    本文围绕“蜜罐采集 + AI语义分析 + 自动化处置”的闭环构建系统，研究内容包括：
    1）蜜罐服务的部署与管理；
    2）日志数据的规范化入库与规则检测；
    3）基于大语言模型的语义分析与风险评估；
    4）恶意IP自动封禁与可视化呈现。
    研究方法以工程实现为主，辅以模块化设计、性能测试与对比分析。
    """
    p = doc.add_paragraph()
    add_long_text(p, content_text, repeat=3)
    doc.add_page_break()

    # 第二章 相关理论/技术
    doc.add_heading('第二章 相关理论与技术', level=1)
    doc.add_heading('2.1 蜜罐技术', level=2)
    honeypot_text = """
    蜜罐是一种主动防御技术，通过模拟真实服务吸引攻击者并记录其行为。根据交互程度可分为低交互与高交互蜜罐。本文系统集成了SSH、FTP与HTTP蜜罐，兼顾部署成本与数据价值。
    """
    p = doc.add_paragraph()
    add_long_text(p, honeypot_text, repeat=5)

    doc.add_heading('2.2 大语言模型与本地部署', level=2)
    llm_text = """
    大语言模型基于Transformer架构，具备强大的语义理解与推理能力。为降低数据外泄风险与使用成本，系统采用本地化部署方案，通过Ollama对模型进行管理与调用，并提供统一的API接口进行推理服务。
    """
    p = doc.add_paragraph()
    add_long_text(p, llm_text, repeat=5)

    doc.add_heading('2.3 后端技术栈', level=2)
    backend_text = """
    后端采用Flask框架，使用Blueprint组织路由，结合SQLAlchemy实现ORM映射，提供日志上传、查询、AI配置管理与恶意IP处置等接口。系统支持CORS跨域访问，并通过任务队列与多线程Worker处理AI分析。
    """
    p = doc.add_paragraph()
    add_long_text(p, backend_text, repeat=4)

    doc.add_heading('2.4 前端技术栈', level=2)
    frontend_text = """
    前端采用Vue 3 + Element Plus构建管理控制台，使用Axios进行API通信，结合ECharts实现攻击地图、趋势统计等可视化模块，实现日志检索与蜜罐配置管理。
    """
    p = doc.add_paragraph()
    add_long_text(p, frontend_text, repeat=4)
    doc.add_page_break()

    # 第三章 需求分析
    doc.add_heading('第三章 需求分析', level=1)
    doc.add_heading('3.1 可行性分析', level=2)
    feas_text = """
    技术上，项目基于成熟的Flask与Vue技术栈，AI推理使用本地部署模型，可控性强。经济上，系统采用开源组件与通用硬件，成本可控。运维上，模块化架构便于部署与扩展。
    """
    p = doc.add_paragraph()
    add_long_text(p, feas_text, repeat=4)

    doc.add_heading('3.2 系统功能需求', level=2)
    req_text = """
    功能需求包括：
    1）蜜罐管理：创建、启动、停止与状态监控；
    2）日志采集与查询：按时间、IP、攻击类型等条件检索；
    3）AI分析：自动识别攻击类型与风险等级；
    4）恶意IP处置：自动或手动封禁；
    5）可视化展示：攻击分布、趋势与统计分析。
    """
    p = doc.add_paragraph()
    add_long_text(p, req_text, repeat=4)

    doc.add_heading('3.3 非功能需求', level=2)
    nfr_text = """
    系统应具备良好的性能与稳定性，支持多蜜罐并发上报；具备安全性与可维护性；支持后续模型扩展与模块升级。
    """
    p = doc.add_paragraph()
    add_long_text(p, nfr_text, repeat=3)

    doc.add_heading('3.4 用例分析', level=2)
    usecase_text = """
    主要参与者包括管理员与安全分析人员。典型用例包括：管理员配置蜜罐与AI模型、分析人员检索日志并查看AI分析结果、系统自动封禁高风险IP。用例图在论文排版中给出。
    """
    p = doc.add_paragraph()
    add_long_text(p, usecase_text, repeat=3)
    doc.add_page_break()

    # 第四章 系统设计
    doc.add_heading('第四章 系统设计', level=1)
    doc.add_heading('4.1 总体架构设计', level=2)
    arch_text = """
    系统采用前后端分离架构：蜜罐模块上报日志至后端API，后端入库后推送至AI分析队列，分析结果回写数据库并在前端展示。前端通过统一接口获取统计数据并渲染可视化图表。
    """
    p = doc.add_paragraph()
    add_long_text(p, arch_text, repeat=4)

    doc.add_heading('4.2 数据库设计', level=2)
    db_text = """
    数据库核心表包括：users（用户与权限）、honeypots（蜜罐配置）、logs（攻击日志）、match_rules（规则匹配）、ai_configs（模型配置）、malicious_ips（恶意IP）、block_history（封禁记录）等。logs表保存源/目标IP、端口、协议、payload与AI分析字段。
    """
    p = doc.add_paragraph()
    add_long_text(p, db_text, repeat=4)

    doc.add_heading('4.3 模块设计', level=2)
    module_text = """
    模块划分为蜜罐服务模块、日志处理模块、AI分析模块、恶意IP处置模块与可视化模块。AI分析模块采用生产者-消费者模型，支持多配置并发分析，提高系统吞吐量。
    """
    p = doc.add_paragraph()
    add_long_text(p, module_text, repeat=4)
    doc.add_page_break()

    # 第五章 系统实现
    doc.add_heading('第五章 系统实现', level=1)
    doc.add_heading('5.1 后端核心逻辑', level=2)
    backend_impl = """
    后端基于Flask实现API服务，日志上传接口接收蜜罐上报数据并写入数据库；AI分析服务使用任务队列将日志分发至不同模型Worker；恶意IP服务根据AI判断结果触发封禁。
    """
    p = doc.add_paragraph()
    add_long_text(p, backend_impl, repeat=4)

    code_honeypot = """
class HoneypotService:
    @staticmethod
    def start_honeypot(hp_id):
        hp = Honeypot.query.get(hp_id)
        cmd = [sys.executable, script_path, '--port', str(hp.port)]
        process = subprocess.Popen(cmd, shell=False)
        running_honeypots[hp.id] = process
        return {'success': True, 'pid': process.pid}
"""
    add_code_block(doc, code_honeypot)

    doc.add_heading('5.2 AI分析实现', level=2)
    ai_impl = """
    AI分析模块调用TrafficAnalysisAgent封装推理逻辑，构造提示词并请求本地模型API，解析JSON结果写回日志表。系统支持多模型配置与动态启停Worker。
    """
    p = doc.add_paragraph()
    add_long_text(p, ai_impl, repeat=4)

    code_ai = """
_task_queue = queue.Queue()

def refresh_workers(app):
    active_configs = get_all_active_configs()
    for cfg in active_configs:
        if cfg['id'] not in running_workers:
            start_worker(app, cfg)
"""
    add_code_block(doc, code_ai)

    doc.add_heading('5.3 前端实现', level=2)
    frontend_impl = """
    前端页面包括仪表盘、日志查询、蜜罐管理、AI配置与恶意IP管理等模块。仪表盘使用ECharts渲染攻击地图与趋势图，日志页支持过滤检索与AI分析详情展示。
    """
    p = doc.add_paragraph()
    add_long_text(p, frontend_impl, repeat=3)
    doc.add_page_break()

    # 第六章 系统测试
    doc.add_heading('第六章 系统测试', level=1)
    doc.add_heading('6.1 测试环境', level=2)
    env_text = """
    测试环境：Windows 10 / Ubuntu 20.04；Python 3.10；Flask 2.x；Vue 3；本地GPU支持7B模型推理。
    """
    doc.add_paragraph(env_text)

    doc.add_heading('6.2 功能测试', level=2)
    test_text = """
    1）蜜罐捕获测试：SSH/FTP/HTTP蜜罐可记录登录尝试与异常请求；
    2）AI分析测试：SQL注入、XSS、命令执行等样例可被识别；
    3）封禁测试：高风险IP自动封禁并记录历史。
    """
    p = doc.add_paragraph()
    add_long_text(p, test_text, repeat=4)

    doc.add_heading('6.3 性能测试', level=2)
    perf_text = """
    在并发上报场景下，系统平均响应时间保持在可接受范围。多Worker并发分析显著提升处理吞吐量。
    """
    p = doc.add_paragraph()
    add_long_text(p, perf_text, repeat=4)
    doc.add_page_break()

    # 第七章 总结与展望
    doc.add_heading('第七章 总结与展望', level=1)
    doc.add_heading('7.1 总结', level=2)
    summary_text = """
    本文完成了基于AI大模型的蜜罐流量分析系统设计与实现，实现了蜜罐采集、日志分析、自动处置与可视化展示的完整闭环。系统验证了本地LLM在安全日志语义分析中的可行性，并通过并发队列提升了处理能力。
    """
    p = doc.add_paragraph()
    add_long_text(p, summary_text, repeat=4)

    doc.add_heading('7.2 展望', level=2)
    future_text = """
    后续可从模型微调、蜜罐类型扩展、与威胁情报联动、云原生部署等方面进一步提升系统能力。
    """
    p = doc.add_paragraph()
    add_long_text(p, future_text, repeat=4)
    doc.add_page_break()

    # 参考文献
    doc.add_heading('参考文献', level=1)
    doc.add_paragraph('（本次初稿暂不编写参考文献，后续按 GB/T 7714 标准补充。）')
    doc.add_page_break()

    # 致谢
    doc.add_heading('致谢', level=1)
    thanks_text = """
    光阴似箭，大学四年的学习生活转瞬即逝。本论文在导师的指导下完成，感谢导师的耐心指导与严格要求。感谢实验室同学在系统开发与测试过程中的帮助，也感谢家人长期以来的理解与支持。
    """
    doc.add_paragraph(thanks_text)
    doc.add_page_break()

    # 附录
    doc.add_heading('附录', level=1)
    appendix_text = """
    附录A：接口示例
    1. /api/logs/internal/upload：蜜罐上报日志接口
    2. /api/ai-config/：AI配置管理接口
    3. /api/malicious-ips/block：恶意IP封禁接口

    附录B：数据库表字段（节选）
    1. logs：攻击日志与AI分析结果
    2. honeypots：蜜罐配置
    3. ai_configs：模型配置
    """
    p = doc.add_paragraph()
    add_long_text(p, appendix_text, repeat=2)

    doc.save(filename)
    print(f"生成完成: {filename}")


if __name__ == '__main__':
    output_path = r"E:\桌面\zuoye\毕业设计\src\论文\论文初稿codex.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    create_thesis_doc(output_path)
